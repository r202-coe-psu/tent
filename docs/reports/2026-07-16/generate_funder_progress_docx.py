#!/usr/bin/env python3
"""Generate funder-facing Smart Shelter progress report (DOCX) — plain language."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path(__file__).with_name(
	'Smart_Shelter_Progress_Report_Funder_2026-07-16.docx'
)

NAVY = RGBColor(0x1A, 0x3A, 0x5C)
MUTED = RGBColor(0x55, 0x55, 0x55)
BLACK = RGBColor(0x22, 0x22, 0x22)
GREEN = RGBColor(0x1B, 0x5E, 0x3B)


def set_run_font(run, *, size=11, bold=False, color=BLACK, name='TH Sarabun New'):
	run.bold = bold
	run.font.size = Pt(size)
	run.font.color.rgb = color
	run.font.name = name
	rPr = run._element.get_or_add_rPr()
	rFonts = rPr.find(qn('w:rFonts'))
	if rFonts is None:
		rFonts = OxmlElement('w:rFonts')
		rPr.insert(0, rFonts)
	rFonts.set(qn('w:ascii'), name)
	rFonts.set(qn('w:hAnsi'), name)
	rFonts.set(qn('w:eastAsia'), name)
	rFonts.set(qn('w:cs'), name)


def add_para(doc, text, *, size=11, bold=False, color=BLACK, align='left', space_after=6, space_before=0):
	p = doc.add_paragraph()
	p.paragraph_format.space_after = Pt(space_after)
	p.paragraph_format.space_before = Pt(space_before)
	p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
	if align == 'center':
		p.alignment = WD_ALIGN_PARAGRAPH.CENTER
	elif align == 'justify':
		p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
	run = p.add_run(text)
	set_run_font(run, size=size, bold=bold, color=color)
	return p


def add_heading(doc, text, level=1):
	sizes = {1: 15, 2: 13, 3: 12}
	return add_para(
		doc,
		text,
		size=sizes.get(level, 12),
		bold=True,
		color=NAVY,
		space_before=14 if level == 1 else 10,
		space_after=6,
	)


def set_cell_shading(cell, hex_color: str):
	tcPr = cell._tc.get_or_add_tcPr()
	shd = OxmlElement('w:shd')
	shd.set(qn('w:fill'), hex_color)
	shd.set(qn('w:val'), 'clear')
	tcPr.append(shd)


def fill_cell(cell, text, *, bold=False, center=False, size=10, color=BLACK):
	cell.text = ''
	p = cell.paragraphs[0]
	if center:
		p.alignment = WD_ALIGN_PARAGRAPH.CENTER
	run = p.add_run(text)
	set_run_font(run, size=size, bold=bold, color=color)


def style_header_row(row, fill='1A3A5C'):
	for cell in row.cells:
		set_cell_shading(cell, fill)
		for p in cell.paragraphs:
			p.alignment = WD_ALIGN_PARAGRAPH.CENTER
			for run in p.runs:
				set_run_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))


def add_table(doc, headers, rows, header_fill='1A3A5C'):
	table = doc.add_table(rows=1 + len(rows), cols=len(headers))
	table.style = 'Table Grid'
	hdr = table.rows[0]
	for i, h in enumerate(headers):
		fill_cell(hdr.cells[i], h, bold=True, center=True, size=10)
	style_header_row(hdr, header_fill)
	for r_i, row in enumerate(rows):
		tr = table.rows[r_i + 1]
		for c_i, val in enumerate(row):
			fill_cell(tr.cells[c_i], str(val), center=(c_i == 0), size=10)
	doc.add_paragraph()
	return table


def add_bullet(doc, text, *, size=11, bold_prefix=None):
	p = doc.add_paragraph(style='List Bullet')
	p.paragraph_format.space_after = Pt(3)
	p.clear()
	if bold_prefix:
		r1 = p.add_run(bold_prefix)
		set_run_font(r1, size=size, bold=True)
		r2 = p.add_run(text)
		set_run_font(r2, size=size)
	else:
		run = p.add_run(text)
		set_run_font(run, size=size)
	return p


def build():
	doc = Document()
	section = doc.sections[0]
	section.top_margin = Cm(2.0)
	section.bottom_margin = Cm(2.0)
	section.left_margin = Cm(2.2)
	section.right_margin = Cm(2.2)

	style = doc.styles['Normal']
	style.font.name = 'TH Sarabun New'
	style.font.size = Pt(11)
	style._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

	# ---- Cover ----
	add_para(doc, 'Smart Shelter — ระบบบริหารศูนย์พักพิง', size=18, bold=True, color=NAVY, align='center', space_after=4)
	add_para(
		doc,
		'รายงานความคืบหน้าการพัฒนาซอฟต์แวร์\nสำหรับเจ้าของทุน / ผู้สนับสนุนโครงการ',
		size=13,
		bold=True,
		align='center',
		space_after=8,
	)
	add_para(doc, 'ประจำวันที่ 18 กรกฎาคม 2026', size=12, color=MUTED, align='center', space_after=14)

	# ---- Executive ----
	add_heading(doc, 'สรุปภาพรวม (อ่านหน้าเดียวนี้ก็เข้าใจสถานะได้)', level=1)
	add_para(
		doc,
		'ความคืบหน้าในรายงานฉบับนี้คิดจากระบบทั้งโครงการ (รวมทุกช่วงงานจนจบโครงการ) '
		'โดยวัดจากความสำเร็จของงานหลักแต่ละชิ้นเฉลี่ยเท่ากัน — ไม่ถ่วงน้ำหนักด้วยจำนวนวันทำงาน (man-days)',
		size=11,
		align='justify',
		space_after=6,
	)
	add_para(
		doc,
		'ความคืบหน้าทั้งระบบ ณ วันนี้: ประมาณ 52%',
		size=13,
		bold=True,
		color=GREEN,
		space_after=4,
	)
	add_para(
		doc,
		'ฐานคิด: งานหลักทั้งหมด 68 งาน · เฉลี่ย % ความสำเร็จของทุกงานเท่ากัน · ครอบคลุมทุก phase จนจบโครงการ · ไม่ใช้ man-days '
		'· ปรับตามสถานะจริงจาก codebase (ค้นหาญาติพร้อมใช้แล้ว, API ส่งข้อมูลกำลังทำ) '
		'ตัวเลขบอร์ดเดิมที่ยังไม่อัปเดตอยู่ที่ 46.96%',
		size=10,
		color=MUTED,
		space_after=6,
	)
	add_para(
		doc,
		'สำคัญ: ระบบทั้งระบบยังไม่ได้ทดสอบการใช้งานจริงครบวงจร '
		'ตัวเลขความคืบหน้าสะท้อนว่ามีฟีเจอร์/โค้ดพัฒนาไปถึงไหนแล้ว '
		'ยังไม่ใช่การยืนยันว่าทุกส่วนพร้อมใช้งานในสถานการณ์จริง',
		size=11,
		bold=True,
		align='justify',
		space_after=6,
	)
	add_para(
		doc,
		'ในช่วงนี้ทีมโฟกัสที่ “ฐานรากของระบบศูนย์พักพิง” — มีโครงลงทะเบียนผู้ประสบภัย จัดโซนพัก '
		'หน้าเว็บสาธารณะ และค้นหาญาติ รวมถึงมีโครงคลังวัสดุและครัวอาหาร '
		'แต่คลังและอาหารยังต้องปรับแก้ requirement และพารามิเตอร์ยังไม่ได้ทดสอบ '
		'งานเฟสถัดไป (เช่น อาสาสมัคร ความปลอดภัยในศูนย์) และ API ส่งข้อมูลยังอยู่ระหว่างดำเนินการ',
		size=11,
		align='justify',
		space_after=6,
	)
	add_para(
		doc,
		'สิ่งที่ยังต้องทำก่อนยืนยันความพร้อมใช้งานจริง:',
		size=11,
		align='justify',
		space_after=4,
	)
	add_bullet(doc, 'ทดสอบทั้งระบบด้วยมือ (manual) เดิน flow การใช้งานจริงครบวงจร — ยังไม่ได้ทำ')
	add_bullet(doc, 'เก็บ feedback จากผู้ใช้งานจริงอย่างเป็นระบบ')
	add_bullet(doc, 'ปรับ requirement ของระบบคลังและครัวอาหาร ให้ตรงการใช้งานจริง และทดสอบพารามิเตอร์ที่เกี่ยวข้อง')

	# Snapshot numbers without internal jargon
	add_heading(doc, '1. สถานะงานแบบภาพรวม (ทั้งระบบ)', level=1)
	add_table(
		doc,
		['ตัวชี้วัด', 'ตัวเลข', 'คำอธิบาย'],
		[
			[
				'ความคืบหน้าทั้งระบบ (ปรับตาม codebase)',
				'≈ 52%',
				'เฉลี่ยความสำเร็จของงานหลักทุกชิ้นเท่ากัน — ยังไม่รวมผลการทดสอบทั้งระบบ',
			],
			['สถานะการทดสอบทั้งระบบ', 'ยังไม่ได้ทดสอบ', 'ยังไม่มีผลการยืนยัน UX / flow ครบวงจรทั้งระบบ'],
			['ตัวเลขอ้างอิงจากบอร์ดเดิม', '46.96%', 'ยังนับค้นหาญาติ/API เป็นยังไม่เริ่ม — ต่ำกว่าสถานะพัฒนาจริง'],
			['จำนวนงานหลักทั้งหมด', '68 งาน', 'ครอบคลุมทุก phase จนจบโครงการ'],
			['มีฟีเจอร์ในระบบแล้ว (ยังรอทดสอบทั้งระบบ)', 'ฐานราก + ค้นหาญาติ + PDPA', 'หน้าเว็บสาธารณะใช้งานค้นหาญาติได้'],
			['กำลังปรับ / กำลังทำ', 'คลัง, อาหาร, API ฯลฯ', 'คลัง–อาหารต้องแก้ requirement และยังไม่ทดสอบพารามิเตอร์'],
			['ยังไม่เริ่ม / พักไว้', 'งานเฟสถัดไป', 'เช่น อาสาสมัคร ความปลอดภัย ส่งมอบ'],
		],
		header_fill='1A3A5C',
	)

	# ---- DONE ----
	add_heading(doc, '2. อะไรที่มีในระบบแล้ว (ยังรอทดสอบทั้งระบบ)', level=1)
	add_para(
		doc,
		'กลุ่มนี้หมายถึง “มีฟีเจอร์/หน้าจอในโปรแกรมแล้ว” จากงานพัฒนา '
		'แต่ระบบทั้งระบบยังไม่ได้ทดสอบการใช้งานจริงครบวงจร '
		'จึงยังไม่นับเป็นพร้อมใช้งานภาคสนาม',
		size=11,
		align='justify',
		space_after=6,
	)

	add_heading(doc, '2.1 การรับผู้ประสบภัยเข้าศูนย์', level=2)
	add_bullet(doc, 'ตั้งค่าข้อมูลศูนย์พักพิง และข้อมูลตั้งต้นของระบบ')
	add_bullet(doc, 'ลงทะเบียนบุคคล / แก้ไขประวัติผู้พักพิง')
	add_bullet(doc, 'คัดกรองเบื้องต้น (กลุ่มเปราะบาง / หมายเหตุสุขภาพ / ช่องทางเร่งด่วน)')
	add_bullet(doc, 'ออกรหัสประจำตัวในศูนย์ และ QR สำหรับเช็คอิน–เช็คเอาต์ระดับบุคคล')
	add_bullet(doc, 'ค้นหาผู้พักพิง สแกน QR และดูประวัติการเข้า–ออก')
	add_bullet(doc, 'แดชบอร์ดภาพรวมศูนย์ฉบับแรก')
	add_bullet(doc, 'จัดการผู้ใช้ระบบหลังบ้าน (ใครเข้าสู่ระบบได้ / บทบาทเบื้องต้น)')
	add_bullet(doc, 'ตั้งค่าข้อมูลหลักที่ใช้ลงทะเบียน (ข้อมูลประชากร / ฟิลด์ครัวเรือน / ข้อมูลประเทศไทย)')
	add_bullet(doc, 'ส่งออกข้อมูลศูนย์เป็นไฟล์ Excel')

	add_heading(doc, '2.2 ครัวเรือนและโซนพักพิง', level=2)
	add_bullet(doc, 'ค้นหาครัวเรือน')
	add_bullet(doc, 'กำหนดโซนพักและความจุของแต่ละโซน')
	add_bullet(doc, 'จัดสรรโซนให้ผู้พัก (ระบบแนะนำได้ แต่ยังเป็นคำเตือน ไม่บังคับ)')
	add_bullet(doc, 'บันทึกสัตว์เลี้ยง สิ่งของ และยานพาหนะที่มากับผู้พัก')
	add_para(
		doc,
		'หมายเหตุ: ไม่มีเช็คอิน–เช็คเอาต์ระดับครัวเรือน (ตัดออกจากขอบเขตแล้ว) — การเข้า–ออกทำที่ระดับบุคคล',
		size=10,
		color=MUTED,
		space_before=2,
		space_after=6,
	)

	add_heading(doc, '2.3 หน้าเว็บสำหรับประชาชน', level=2)
	add_bullet(doc, 'หน้าแรกของพอร์ทัลสาธารณะ พร้อมตัวเลขสถานการณ์แบบใกล้เรียลไทม์')
	add_bullet(doc, 'หน้าดูภาพรวมศูนย์พักพิงสำหรับประชาชน')
	add_bullet(doc, 'ค้นหาญาติ/ครอบครัวแบบสาธารณะ — ทดสอบใช้งานได้แล้วในส่วนนี้')
	add_bullet(doc, 'คุ้มครองข้อมูลส่วนบุคคล (PDPA): ปกปิดข้อมูลบางส่วน และรองรับการไม่แสดงในผลการค้นหาตามความประสงค์')
	add_bullet(doc, 'ผู้บริจาคแจ้งของล่วงหน้าได้ / หน้าเว็บจองคิวส่งของ (ส่วนพอร์ทัล)')

	add_para(
		doc,
		'สรุปสั้นๆ ของส่วนที่มีในระบบแล้ว: “รับคนเข้าศูนย์ → จัดโซน → ให้ประชาชนดูข้อมูลศูนย์และค้นหาญาติ” '
		'มีโครงใช้งานได้ แต่ยังต้องทดสอบทั้งระบบร่วมกับโมดูลอื่น',
		size=11,
		bold=True,
		align='justify',
		space_before=4,
		space_after=8,
	)

	# ---- IN PROGRESS ----
	add_heading(doc, '3. อะไรที่กำลังดำเนินการอยู่ / ยังต้องปรับ', level=1)
	add_para(
		doc,
		'กลุ่มนี้มีโครงระบบและหน้าจอบางส่วนแล้ว แต่ยังไม่ครบ ยังต้องปรับ requirement '
		'หรือยังไม่ได้ทดสอบพารามิเตอร์/การใช้งานจริง',
		size=11,
		align='justify',
		space_after=6,
	)
	add_table(
		doc,
		['ฟีเจอร์', 'สถานะโดยประมาณ', 'กำลังทำอะไรอยู่ / ประเด็นค้าง'],
		[
			[
				'ระบบคลังวัสดุ (รับเข้า / จ่ายออก / โอนข้ามศูนย์)',
				'มีโครงแล้ว — ยังไม่ปิด',
				'ยังต้องปรับแก้ requirement และพารามิเตอร์ที่เกี่ยวข้องยังไม่ได้ทดสอบ',
			],
			[
				'ระบบครัวอาหาร (แผนมื้อ / เบิกวัตถุดิบ / บันทึกแจกอาหาร)',
				'มีโครงแล้ว — ยังไม่ปิด',
				'ยังต้องปรับแก้ requirement และพารามิเตอร์ที่เกี่ยวข้องยังไม่ได้ทดสอบ',
			],
			[
				'แค็ตตาล็อกสินค้าในคลัง / แดชบอร์ดสต็อกและจุดสั่งของใหม่',
				'ราวครึ่งทาง',
				'กำลังทำหน้าจอและกฎการแจ้งเตือนเมื่อของใกล้หมด',
			],
			[
				'สร้างครัวเรือน ผูกสมาชิก และกำหนดหัวหน้าครัวเรือน',
				'ราว 2 ใน 3',
				'ส่วนใหญ่มีโครงแล้ว ยังเหลือจุดเชื่อม/เคสพิเศษบางส่วน',
			],
			[
				'เครื่องคำนวณทรัพยากรรายวัน (อาหาร/ของจำเป็นจากจำนวนผู้พัก)',
				'เกือบเสร็จด้านโค้ด',
				'แกนคำนวณใกล้จบ แต่ยังผูกกับพารามิเตอร์ที่ยังไม่ได้ทดสอบ',
			],
			[
				'ตั้งค่าอัตราส่วนมาตรฐาน (SOP) สำหรับคำนวณของ',
				'เกือบเสร็จด้านโค้ด',
				'หน้าตั้งค่าส่วนใหญ่มีแล้ว ยังต้องทดสอบพารามิเตอร์กับการใช้งานจริง',
			],
			[
				'ระบบจองคิวบริจาค / ตัดรอบรับบริจาค / เส้นทางตรวจสอบการรับของ',
				'ราว 1 ใน 5 ถึงครึ่งหนึ่ง',
				'โครงมีแล้ว แต่ยังไม่ครบวงจรการใช้งานจริง',
			],
			[
				'สิทธิ์ผู้ใช้ละเอียดขึ้น (ใครเห็น/แก้ข้อมูลส่วนไหนได้)',
				'ราวครึ่งทาง',
				'กำลังขยายบทบาทและสิทธิ์ระดับฟิลด์',
			],
			[
				'การทำงานแบบ remote-first (ศูนย์กลางก่อน แล้วค่อยสลับไปเครื่องในศูนย์เมื่อเน็ตล่ม)',
				'ราวครึ่งทาง',
				'กำลังทำให้ระบบเลือกเส้นทางเชื่อมต่อให้ถูกต้องและปลอดภัย',
			],
			[
				'ส่งต่อเคสไปหน่วยงานภายนอก (referral)',
				'ราวไม่ถึงครึ่ง',
				'มีโครงงานแล้ว ยังไม่จบ flow การส่งต่อ',
			],
			[
				'คำถามที่พบบ่อย (FAQ) บนเว็บสาธารณะ + หน้าตั้งค่า',
				'ราวครึ่งทาง',
				'กำลังทำเนื้อหาแบบอัปเดตได้จากระบบ',
			],
			[
				'API ส่งข้อมูลภาพรวมไประบบวิเคราะห์ / ศูนย์บัญชาการ',
				'กำลังดำเนินการ',
				'มี worker ซิงก์ข้อมูลจาก CouchDB → MongoDB แล้ว (อยู่บน branch งานนี้) ยังไม่ปิดงานทั้งหมด',
			],
		],
	)

	# ---- NOT DONE ----
	add_heading(doc, '4. อะไรที่ยังไม่เสร็จ / ยังไม่เริ่ม', level=1)

	add_heading(doc, '4.1 การทดสอบทั้งระบบ — ยังไม่ได้ทำ', level=2)
	add_para(
		doc,
		'ณ ขณะนี้ระบบทั้งระบบยังไม่ได้ผ่านการทดสอบการใช้งานจริงครบวงจร '
		'รวมถึง flow หลักบนหน้างานและจุดเชื่อมระหว่างโมดูล',
		size=11,
		align='justify',
		space_after=4,
	)
	add_bullet(
		doc,
		'ยังไม่ได้เดินทดสอบด้วยมือครบ flow เช่น ลงทะเบียน → คัดกรอง → ออก QR → เช็คอินบุคคล → จัดโซน → คลัง/ครัว',
		bold_prefix='ทดสอบทั้งระบบ (manual): ',
	)
	add_bullet(
		doc,
		'ยังไม่มีวงจรรับ feedback จากเจ้าหน้าที่ศูนย์อย่างเป็นระบบ แล้วนำกลับมาปรับหน้าจอ',
		bold_prefix='Feedback จากผู้ใช้จริง: ',
	)
	add_bullet(
		doc,
		'พารามิเตอร์ของคลังและอาหารยังไม่ได้ทดสอบ และ requirement ยังต้องปรับแก้',
		bold_prefix='คลังและอาหาร: ',
	)

	add_heading(doc, '4.2 งานที่ยังไม่เริ่ม (เฟสถัดไป / พักไว้ก่อน)', level=2)
	add_bullet(doc, 'ระบบอาสาสมัคร: ลงทะเบียน ทักษะ ความพร้อม และจัดเวร/จับคู่งาน')
	add_bullet(doc, 'ระบบรับเรื่องและติดตามเหตุการณ์ความปลอดภัยในศูนย์')
	add_bullet(doc, 'จำลองสถานการณ์ what-if ของการจัดสรรทรัพยากร')
	add_bullet(doc, 'คู่มือผู้ใช้ / คู่มือปฏิบัติการ และการฝึกอบรมส่งมอบ')
	add_bullet(doc, 'ระบบติดตั้งอัตโนมัติ (deployment) และชุดส่งมอบปิดโครงการ')
	add_bullet(doc, 'งานตกแต่ง/เก็บรายละเอียดบางส่วน และการทดสอบรับมอบข้ามโมดูล')
	add_para(
		doc,
		'หมายเหตุ: ค้นหาญาติ (+ PDPA) อยู่ในกลุ่มที่มีในระบบแล้ว · API ส่งข้อมูลอยู่ในกลุ่มกำลังดำเนินการ · '
		'เช็คอิน–เช็คเอาต์ระดับครัวเรือนตัดออกจากขอบเขตแล้ว',
		size=10,
		color=MUTED,
		space_before=4,
		space_after=6,
	)

	# ---- By area overview ----
	add_heading(doc, '5. ภาพรวมแยกตามกลุ่มงาน', level=1)
	add_para(
		doc,
		'ตารางนี้ช่วยมองว่ากลุ่มไหนเดินหน้าแล้ว และกลุ่มไหนยังอยู่ข้างหน้า '
		'(ตัวเลขเป็นความคืบหน้าด้านการพัฒนา — ไม่ใช่ผลการทดสอบทั้งระบบ)',
		size=11,
		align='justify',
		space_after=6,
	)
	add_table(
		doc,
		['กลุ่มงาน', 'ความคืบหน้า', 'สรุปสั้นๆ'],
		[
			['ทะเบียนผู้พักพิง / ตั้งค่าศูนย์', '70%', 'มีฟีเจอร์ในระบบแล้ว — ยังรอทดสอบทั้งระบบ'],
			['ครัวเรือนและโซนพัก', '70%', 'มีค้นหาครัวเรือน/โซน — ไม่มีเช็คอิน–เอาต์ระดับครัวเรือน'],
			['คลังวัสดุ', '55%', 'มีโครงแล้ว; ต้องแก้ requirement และพารามิเตอร์ยังไม่ได้ทดสอบ'],
			['ครัวอาหาร', '55%', 'มีโครงแล้ว; ต้องแก้ requirement และพารามิเตอร์ยังไม่ได้ทดสอบ'],
			['หน้าเว็บประชาชน + ค้นหาญาติ (PDPA)', '85%', 'ค้นหาญาติทดสอบใช้งานได้แล้วในพอร์ทัล'],
			['ระบบบริจาค', '40%', 'รับบริจาคเบื้องต้นมีโครง; วงจรเต็มยังไม่จบ'],
			['คำนวณทรัพยากรตามมาตรฐาน', '50%', 'โค้ดใกล้เสร็จ; พารามิเตอร์ยังต้องทดสอบ'],
			['API ส่งข้อมูล / ศูนย์บัญชาการ', '45%', 'worker ซิงก์ CouchDB→MongoDB อยู่ระหว่างพัฒนา'],
			['โครงสร้างระบบส่วนกลาง', '14%', 'สิทธิ์/การเชื่อมต่อกำลังทำ; ดีพลอยและ UAT ใหญ่ยังไม่เริ่ม'],
			['อาสาสมัคร / ความปลอดภัยในศูนย์', '0%', 'ยังไม่เริ่ม — อยู่ในแผนเฟสถัดไป'],
			['การทดสอบทั้งระบบ', '0%', 'ยังไม่ได้ทดสอบทั้งระบบครบวงจร'],
		],
	)

	# ---- Next ----
	add_heading(doc, '6. สิ่งที่จะทำต่อจากนี้', level=1)
	add_bullet(doc, 'ทดสอบทั้งระบบแบบเดิน flow ด้วยมือ และจดจุดที่ติดขัด')
	add_bullet(doc, 'ปรับ requirement ของระบบคลังและครัวอาหาร ให้ตรงการใช้งานจริง')
	add_bullet(doc, 'ทดสอบพารามิเตอร์ของคลัง อาหาร และส่วนคำนวณที่เกี่ยวข้อง')
	add_bullet(doc, 'เก็บ feedback จากผู้ทดลอง แล้วปรับหน้าจอ/ลำดับงานให้ใช้งานง่ายขึ้น')
	add_bullet(doc, 'เก็บงานที่ค้างในกลุ่ม “กำลังทำ” และงานเฟสถัดไปตามแผน')

	# ---- Close ----
	add_heading(doc, '7. สรุป', level=1)
	add_para(
		doc,
		'ความคืบหน้าทั้งระบบ ณ วันที่ 18 กรกฎาคม 2026 อยู่ที่ประมาณ 52% '
		'โดยวัดจากความสำเร็จของงานหลักทุกชิ้นเฉลี่ยเท่ากัน ครอบคลุมทุก phase จนจบโครงการ '
		'ไม่ใช้ man-days และปรับตามสถานะจริงจาก codebase '
		'(ค้นหาญาติพร้อมใช้ในพอร์ทัลรวม PDPA, API ส่งข้อมูลกำลังดำเนินการ)',
		size=11,
		align='justify',
		space_after=8,
	)
	add_para(
		doc,
		'อย่างไรก็ตาม ระบบทั้งระบบยังไม่ได้ทดสอบการใช้งานจริงครบวงจร '
		'ระบบคลังและครัวอาหารยังต้องปรับแก้ requirement และพารามิเตอร์ยังไม่ได้ทดสอบ '
		'เช็คอิน–เช็คเอาต์ระดับครัวเรือนตัดออกจากขอบเขตแล้ว '
		'ไทม์ไลน์ยังมุ่ง go-live ช่วงต้นเดือนกันยายน 2026',
		size=11,
		align='justify',
		space_after=16,
	)

	add_para(
		doc,
		'— จบรายงาน —',
		size=10,
		color=MUTED,
		align='center',
		space_before=10,
	)
	add_para(
		doc,
		'จัดทำเพื่อสื่อสารกับเจ้าของทุน · ข้อมูล ณ วันที่ 18 กรกฎาคม 2026 · ความคืบหน้าทั้งระบบ ≈ 52% '
		'(วัดจากความสำเร็จของงาน + ปรับตาม codebase · ทั้งระบบยังไม่ได้ทดสอบ)',
		size=9,
		color=MUTED,
		align='center',
	)

	doc.save(OUT)
	print(f'Wrote {OUT}')


if __name__ == '__main__':
	build()
